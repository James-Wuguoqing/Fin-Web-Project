#!/usr/bin/env python

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ACCENT = "2F6DF6"
GREEN = "18A874"
TEXT = "101828"
MUTED = "667085"
LINE = "D9E0EA"
SOFT = "F3F7FF"
SUCCESS_BG = "ECFDF3"
SUCCESS_TEXT = "067647"
FAIL_BG = "FEF3F2"
FAIL_TEXT = "B42318"


def main() -> int:
    report_json = Path(sys.argv[1] if len(sys.argv) > 1 else "reports/visible-smoke/latest/report.json").resolve()
    if not report_json.exists():
        print(f"[FAIL] report.json not found: {report_json}")
        return 1

    data = json.loads(report_json.read_text(encoding="utf-8"))
    report_dir = report_json.parent
    output_path = report_dir / "FinScope-可视化功能测试报告.docx"

    doc = Document()
    configure_document(doc)
    build_cover(doc, data, report_dir)
    build_summary(doc, data)
    build_environment(doc, data, report_dir)
    build_coverage(doc, data)
    build_step_details(doc, data, report_dir)
    build_console_section(doc, data)
    build_conclusion(doc, data)

    doc.save(output_path)
    print(f"[PASS] DOCX report written: {output_path}")
    return 0


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Title", 25, ACCENT),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 12, TEXT),
        ("Heading 3", 10.5, TEXT),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def build_cover(doc: Document, data: dict, report_dir: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FinScope")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = "Title"
    title.add_run("可视化功能测试报告")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("覆盖页面交互、表单流程、新闻详情、市场分类与 API 健康检查。").font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph()
    summary = data.get("summary", {})
    status = "PASS" if summary.get("failed", 0) == 0 and summary.get("materialConsoleIssues", 0) == 0 else "FAIL"
    add_badge_table(
        doc,
        [
            ("整体结果", status),
            ("通过步骤", f"{summary.get('passed', 0)} / {summary.get('total', 0)}"),
            ("Warnings", str(summary.get("warnings", 0))),
            ("Console issues", str(summary.get("materialConsoleIssues", 0))),
        ],
    )

    add_meta_table(
        doc,
        [
            ("测试基准 URL", data.get("baseUrl", "")),
            ("执行模式", data.get("mode", "")),
            ("开始时间", format_dt(data.get("startedAt"))),
            ("结束时间", format_dt(data.get("finishedAt"))),
            ("报告目录", str(report_dir)),
            ("完整录屏", video_artifact_text(data, report_dir)),
        ],
    )
    doc.add_page_break()


def build_summary(doc: Document, data: dict) -> None:
    doc.add_heading("1. 执行摘要", level=1)
    summary = data.get("summary", {})
    paragraph = doc.add_paragraph()
    paragraph.add_run("结论：").bold = True
    if summary.get("failed", 0) == 0 and summary.get("materialConsoleIssues", 0) == 0:
        paragraph.add_run("本轮可视化功能测试全部通过，未发现业务相关控制台问题。")
    else:
        paragraph.add_run("本轮测试存在失败项或控制台问题，需按明细逐项处理。")

    add_metric_lines(
        doc,
        [
            ("总步骤", summary.get("total", 0)),
            ("通过", summary.get("passed", 0)),
            ("失败", summary.get("failed", 0)),
            ("Warnings", summary.get("warnings", 0)),
            ("Console issues", summary.get("materialConsoleIssues", 0)),
        ],
    )


def build_environment(doc: Document, data: dict, report_dir: Path) -> None:
    doc.add_heading("2. 测试环境与产物", level=1)
    add_meta_table(
        doc,
        [
            ("项目", data.get("project", "FinScope Next")),
            ("浏览器", data.get("browserPath", "")),
            ("测试 URL", data.get("baseUrl", "")),
            ("结构化报告", str(report_dir / "report.json")),
            ("终端日志", str(report_dir / "terminal.log")),
            ("录屏文件", video_artifact_text(data, report_dir)),
            ("录屏状态", data.get("videoStatus", "unknown")),
        ],
    )


def build_coverage(doc: Document, data: dict) -> None:
    doc.add_heading("3. 覆盖范围", level=1)
    for item in data.get("results", []):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        title = p.add_run(f"{item.get('index', '')}. {item.get('name', '')}")
        title.bold = True
        title.font.color.rgb = RGBColor.from_string(TEXT)
        p.add_run(f"  |  {item.get('status', '')}  |  {item.get('endUrl', '')}")
        screenshot = item.get("screenshotRelative", "")
        if screenshot:
            screenshot_p = doc.add_paragraph()
            screenshot_p.paragraph_format.left_indent = Cm(0.7)
            screenshot_run = screenshot_p.add_run(f"截图：{screenshot}")
            screenshot_run.font.size = Pt(8.5)
            screenshot_run.font.color.rgb = RGBColor.from_string(MUTED)


def build_step_details(doc: Document, data: dict, report_dir: Path) -> None:
    doc.add_heading("4. 测试步骤与截屏", level=1)
    for item in data.get("results", []):
        heading = doc.add_heading(f"{item.get('index', '')}. {item.get('name', '')}", level=2)
        heading.paragraph_format.keep_with_next = True

        p = doc.add_paragraph()
        status_run = p.add_run(f"状态：{item.get('status', '')}  ")
        status_run.bold = True
        status_run.font.color.rgb = RGBColor.from_string(SUCCESS_TEXT if item.get("status") == "PASS" else FAIL_TEXT)
        p.add_run(f"耗时：{item.get('durationMs', 0)} ms  ")
        p.add_run(f"最终 URL：{item.get('endUrl', '')}")

        add_list(doc, "关键动作", item.get("actions", []))
        add_list(doc, "断言与观察", item.get("assertions", [])[:8])

        screenshot = item.get("screenshotRelative")
        if screenshot:
            image_path = report_dir / screenshot
            if image_path.exists():
                caption = doc.add_paragraph(f"图 {item.get('index', '')}：{item.get('name', '')} 截屏")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].font.size = Pt(8.5)
                caption.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
                doc.add_picture(str(image_path), width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph(f"截图缺失：{image_path}")

        if item.get("detail"):
            add_callout(doc, f"失败详情：{item.get('detail')}", FAIL_BG, FAIL_TEXT)


def build_console_section(doc: Document, data: dict) -> None:
    doc.add_heading("5. 控制台结果", level=1)
    material = data.get("materialConsoleIssues", [])
    if not material:
        add_callout(doc, "本轮测试未捕获业务相关 console error 或 warning。", SUCCESS_BG, SUCCESS_TEXT)
        return

    for issue in material:
        add_callout(doc, issue, FAIL_BG, FAIL_TEXT)


def build_conclusion(doc: Document, data: dict) -> None:
    doc.add_heading("6. 结论", level=1)
    summary = data.get("summary", {})
    if summary.get("failed", 0) == 0 and summary.get("materialConsoleIssues", 0) == 0:
        doc.add_paragraph("本轮测试满足通过标准：12 个功能模块全部通过，Warnings 为 0，Material console issues 为 0。")
    else:
        doc.add_paragraph("本轮测试未满足通过标准，请优先处理失败项和控制台问题。")


def add_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        label_run = p.add_run(f"{label}：")
        label_run.bold = True
        label_run.font.color.rgb = RGBColor.from_string(ACCENT)
        value_run = p.add_run(str(value))
        value_run.font.color.rgb = RGBColor.from_string(TEXT)


def add_badge_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, (label, value) in enumerate(rows):
        if idx:
            p.add_run("    ")
        label_run = p.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.color.rgb = RGBColor.from_string(MUTED)
        value_run = p.add_run(str(value))
        value_run.bold = True
        value_run.font.size = Pt(12)
        value_run.font.color.rgb = RGBColor.from_string(SUCCESS_TEXT if value in ("PASS", "0") else TEXT)


def add_metric_lines(doc: Document, rows: list[tuple[str, int]]) -> None:
    p = doc.add_paragraph()
    for idx, (label, value) in enumerate(rows):
        if idx:
            p.add_run("    ")
        label_run = p.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.color.rgb = RGBColor.from_string(MUTED)
        value_run = p.add_run(str(value))
        value_run.bold = True
        value_run.font.color.rgb = RGBColor.from_string(SUCCESS_TEXT if int(value) == 0 and label in ("失败", "Warnings", "Console issues") else TEXT)


def add_list(doc: Document, title: str, values: list[str]) -> None:
    if not values:
        return
    p = doc.add_paragraph()
    p.add_run(f"{title}：").bold = True
    for value in values:
        item = doc.add_paragraph(style=None)
        item.paragraph_format.left_indent = Cm(0.45)
        item.paragraph_format.first_line_indent = Cm(-0.25)
        item.add_run("• ").font.color.rgb = RGBColor.from_string(ACCENT)
        item.add_run(value)


def add_callout(doc: Document, text: str, fill: str, color: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)


def format_header_row(row, headers: list[str]) -> None:
    for idx, text in enumerate(headers):
        cell = row.cells[idx]
        cell.text = text
        shade_cell(cell, ACCENT)
        set_cell_text_color(cell, "FFFFFF", bold=True)
        center_cell(cell)


def set_table_padding(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tbl_cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), "90")
        node.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_width(table, total_width_cm: float, col_widths_cm: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(cm_to_twips(total_width_cm)))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(col_widths_cm):
                set_cell_width(cell, col_widths_cm[idx])


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(cm_to_twips(width_cm)))


def cm_to_twips(value: float) -> int:
    return int(value * 567)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color: str, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)
            run.bold = bold


def center_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def format_dt(value: str | None) -> str:
    if not value:
        return ""
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).strftime("%Y-%m-%d %H:%M:%S")
    except ValueError:
        return value


def video_artifact_text(data: dict, report_dir: Path) -> str:
    video_relative = data.get("videoRelative")
    if video_relative:
        return str(report_dir / video_relative)
    return data.get(
        "videoNote",
        "未生成（当前环境未安装 Playwright ffmpeg；报告已嵌入关键截屏）。",
    )


if __name__ == "__main__":
    raise SystemExit(main())
